from __future__ import annotations

import io
from typing import Any
from zipfile import ZIP_DEFLATED, ZipFile

import pymupdf
from docx import Document as WordDocument
from httpx import AsyncClient

from tests.phase3_helpers import create_account, create_workspace

PDF_MIME = "application/pdf"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def make_pdf(*pages: str | None, draw_image_shape: bool = False) -> bytes:
    pdf = pymupdf.open()
    for text in pages or (None,):
        page = pdf.new_page()
        if text:
            page.insert_text((72, 72), text, fontsize=11)
        if draw_image_shape:
            page.draw_rect(pymupdf.Rect(72, 90, 180, 150), color=(0, 0, 0), fill=(0.5, 0.5, 0.5))
    content = pdf.tobytes()
    pdf.close()
    return content


def make_scanned_pdf(*pages: str) -> bytes:
    """Create a small deterministic image-only PDF without external fonts."""

    scanned = pymupdf.open()
    for text in pages:
        source = pymupdf.open()
        source_page = source.new_page()
        source_page.insert_textbox(
            pymupdf.Rect(72, 100, 523, 742),
            text,
            fontsize=24,
            lineheight=1.4,
        )
        pixmap = source_page.get_pixmap(dpi=200, colorspace=pymupdf.csRGB, alpha=False)
        target_page = scanned.new_page(width=source_page.rect.width, height=source_page.rect.height)
        target_page.insert_image(target_page.rect, stream=pixmap.tobytes("png"))
        source.close()
    content = scanned.tobytes(garbage=4, deflate=True)
    scanned.close()
    return content


def make_mixed_pdf(direct_text: str, scanned_text: str) -> bytes:
    mixed = pymupdf.open(stream=make_pdf(direct_text), filetype="pdf")
    scanned = pymupdf.open(stream=make_scanned_pdf(scanned_text), filetype="pdf")
    mixed.insert_pdf(scanned)
    content = mixed.tobytes()
    scanned.close()
    mixed.close()
    return content


def make_docx(
    *,
    heading: str | None = None,
    paragraphs: tuple[str, ...] = (),
    table: tuple[tuple[str, ...], ...] = (),
) -> bytes:
    document = WordDocument()
    if heading is not None:
        document.add_heading(heading, level=1)
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    if table:
        word_table = document.add_table(rows=len(table), cols=len(table[0]))
        for row_index, row in enumerate(table):
            for column_index, value in enumerate(row):
                word_table.cell(row_index, column_index).text = value
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def make_malformed_docx() -> bytes:
    buffer = io.BytesIO()
    with ZipFile(buffer, "w", ZIP_DEFLATED) as archive:
        archive.writestr(
            "[Content_Types].xml",
            """<?xml version="1.0"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
<Override PartName="/word/document.xml"
 ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>""",
        )
        archive.writestr("word/document.xml", "<not-valid-word-xml")
    return buffer.getvalue()


async def create_case(
    client: AsyncClient,
    workspace_id: str,
    headers: dict[str, str],
    name: str = "Extraction case",
) -> dict[str, Any]:
    response = await client.post(
        f"/api/v1/workspaces/{workspace_id}/cases",
        headers=headers,
        json={"name": name},
    )
    assert response.status_code == 201, response.text
    return response.json()


async def create_document_context(
    client: AsyncClient,
    email: str,
    *,
    filename: str = "agreement.pdf",
    content: bytes | None = None,
    mime_type: str = PDF_MIME,
) -> tuple[
    dict[str, Any],
    dict[str, Any],
    dict[str, Any],
    dict[str, Any],
    dict[str, str],
    bytes,
]:
    user, headers = await create_account(client, email)
    workspace = await create_workspace(client, headers)
    legal_case = await create_case(client, workspace["id"], headers)
    original = content if content is not None else make_pdf(
        "This Agreement contains sufficient selectable legal text for direct extraction."
    )
    response = await client.post(
        f"/api/v1/workspaces/{workspace['id']}/cases/{legal_case['id']}/documents",
        headers=headers,
        files={"file": (filename, original, mime_type)},
    )
    assert response.status_code == 201, response.text
    return user, workspace, legal_case, response.json(), headers, original


def document_base(workspace: dict[str, Any], legal_case: dict[str, Any]) -> str:
    return f"/api/v1/workspaces/{workspace['id']}/cases/{legal_case['id']}/documents"
